import argparse
import json
import os
import sys
from xml.sax.saxutils import escape


def load_json(path):
    with open(path, "r", encoding="utf-8") as handle:
        return json.load(handle)


def build_contact_line(contact):
    if not contact:
        return ""
    parts = []
    for key in ["email", "phone", "website", "linkedin", "address"]:
        value = contact.get(key)
        if value:
            parts.append(value)
    for item in contact.get("other", []) or []:
        if item:
            parts.append(item)
    return " | ".join(parts)


def format_education(edu):
    lines = []
    school = edu.get("school", "")
    location = edu.get("location", "")
    if school or location:
        lines.append(" - ".join([p for p in [school, location] if p]))
    degree = edu.get("degree", "")
    major = edu.get("major", "")
    minor = edu.get("minor", "")
    degree_parts = [degree]
    if major:
        degree_parts.append("Major: " + major)
    if minor:
        degree_parts.append("Minor: " + minor)
    degree_line = "; ".join([p for p in degree_parts if p])
    if degree_line:
        lines.append(degree_line)
    dates = edu.get("dates", "")
    if dates:
        lines.append(dates)
    gpa = edu.get("gpa", "")
    awards = edu.get("awards", []) or []
    coursework = edu.get("coursework", []) or []
    extras = []
    if gpa:
        extras.append("GPA: " + gpa)
    if awards:
        extras.append("Awards: " + ", ".join(awards))
    if coursework:
        extras.append("Coursework: " + ", ".join(coursework))
    if extras:
        lines.append("; ".join(extras))
    return lines


def resolve_font_file(font_name, font_file):
    if font_file:
        return font_file
    if not font_name:
        return None
    windir = os.environ.get("WINDIR", r"C:\Windows")
    font_dir = os.path.join(windir, "Fonts")
    if not os.path.isdir(font_dir):
        return None

    def normalize(text):
        return "".join(ch.lower() for ch in text if ch.isalnum())

    target = normalize(font_name)
    matches = []
    for filename in os.listdir(font_dir):
        lower = filename.lower()
        if not (lower.endswith(".ttf") or lower.endswith(".otf")):
            continue
        if target in normalize(filename):
            matches.append(os.path.join(font_dir, filename))
    for path in matches:
        if path.lower().endswith(".ttf"):
            return path
    return matches[0] if matches else None


def resolve_pdf_font_name(font_name):
    if not font_name:
        return "Helvetica"
    mapping = {
        "helvetica": "Helvetica",
        "timesnewroman": "Times-Roman",
        "times new roman": "Times-Roman",
        "georgia": "Times-Roman",
        "cambria": "Times-Roman",
        "garamond": "Times-Roman",
        "arial": "Helvetica",
        "calibri": "Helvetica",
        "verdana": "Helvetica",
        "roboto": "Helvetica",
        "lato": "Helvetica",
        "opensans": "Helvetica",
        "open sans": "Helvetica",
    }
    return mapping.get(font_name.lower(), "Helvetica")


def create_docx(data, output_path, font_name):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        print("Missing dependency: python-docx. Run: pip install python-docx")
        sys.exit(1)

    document = Document()
    style = document.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)

    name = data.get("name", "")
    if name:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(name)
        run.font.name = font_name
        run.font.size = Pt(18)
        run.bold = True

    contact_line = build_contact_line(data.get("contact", {}))
    if contact_line:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(contact_line)
        run.font.name = font_name
        run.font.size = Pt(10)

    def add_heading(text):
        if not text:
            return
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text.upper())
        run.bold = True
        run.font.name = font_name
        run.font.size = Pt(12)

    def add_entry_header(parts):
        line = " | ".join([p for p in parts if p])
        if not line:
            return
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = font_name
        run.bold = True

    def add_bullets(bullets):
        for bullet in bullets or []:
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(bullet)
            run.font.name = font_name
            run.font.size = Pt(11)

    education = data.get("education", [])
    if education:
        add_heading("Education")
        for edu in education:
            for line in format_education(edu):
                paragraph = document.add_paragraph(line)
                paragraph.runs[0].font.name = font_name

    professional = data.get("professional_experience", [])
    if professional:
        add_heading("Professional Experience")
        for item in professional:
            header_parts = [
                item.get("title", ""),
                item.get("company", ""),
                item.get("location", ""),
                item.get("dates", ""),
            ]
            add_entry_header(header_parts)
            add_bullets(item.get("bullets", []))

    leadership = data.get("leadership_experience", [])
    if leadership:
        add_heading("Leadership Experience")
        for item in leadership:
            header_parts = [
                item.get("title", ""),
                item.get("company", ""),
                item.get("location", ""),
                item.get("dates", ""),
            ]
            add_entry_header(header_parts)
            add_bullets(item.get("bullets", []))

    skills = data.get("skills", {})
    if skills:
        add_heading("Skills")
        technical = skills.get("technical", []) or []
        languages = skills.get("languages", []) or []
        other = skills.get("other", []) or []
        if technical:
            document.add_paragraph("Technical Skills: " + ", ".join(technical))
        if languages:
            document.add_paragraph("Languages: " + ", ".join(languages))
        if other:
            document.add_paragraph(", ".join(other))

    document.save(output_path)


def create_pdf(data, output_path, font_name, font_file):
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.units import inch
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError:
        print("Missing dependency: reportlab. Run: pip install reportlab")
        sys.exit(1)

    pdf_font = None
    font_path = resolve_font_file(font_name, font_file)
    if font_path and font_path.lower().endswith(".ttf"):
        pdf_font = font_name
        try:
            pdfmetrics.registerFont(TTFont(pdf_font, font_path))
        except Exception:
            pdf_font = None

    if not pdf_font:
        pdf_font = resolve_pdf_font_name(font_name)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontName=pdf_font,
        fontSize=18,
        spaceAfter=6,
    )
    contact_style = ParagraphStyle(
        "Contact",
        parent=styles["Normal"],
        fontName=pdf_font,
        fontSize=10,
        spaceAfter=10,
    )
    heading_style = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontName=pdf_font,
        fontSize=12,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName=pdf_font,
        fontSize=11,
        spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "Bullet",
        parent=styles["Normal"],
        fontName=pdf_font,
        fontSize=11,
        leftIndent=14,
        bulletIndent=0,
        spaceAfter=1,
    )

    story = []
    name = data.get("name", "")
    if name:
        story.append(Paragraph(escape(name), title_style))

    contact_line = build_contact_line(data.get("contact", {}))
    if contact_line:
        story.append(Paragraph(escape(contact_line), contact_style))

    def add_heading(text):
        if text:
            story.append(Paragraph(escape(text.upper()), heading_style))

    def add_entry_header(parts):
        line = " | ".join([p for p in parts if p])
        if line:
            story.append(Paragraph(escape(line), body_style))

    def add_bullets(bullets):
        for bullet in bullets or []:
            story.append(Paragraph(escape(bullet), bullet_style, bulletText="-"))

    education = data.get("education", [])
    if education:
        add_heading("Education")
        for edu in education:
            for line in format_education(edu):
                story.append(Paragraph(escape(line), body_style))
            story.append(Spacer(1, 6))

    professional = data.get("professional_experience", [])
    if professional:
        add_heading("Professional Experience")
        for item in professional:
            header_parts = [
                item.get("title", ""),
                item.get("company", ""),
                item.get("location", ""),
                item.get("dates", ""),
            ]
            add_entry_header(header_parts)
            add_bullets(item.get("bullets", []))
            story.append(Spacer(1, 6))

    leadership = data.get("leadership_experience", [])
    if leadership:
        add_heading("Leadership Experience")
        for item in leadership:
            header_parts = [
                item.get("title", ""),
                item.get("company", ""),
                item.get("location", ""),
                item.get("dates", ""),
            ]
            add_entry_header(header_parts)
            add_bullets(item.get("bullets", []))
            story.append(Spacer(1, 6))

    skills = data.get("skills", {})
    if skills:
        add_heading("Skills")
        technical = skills.get("technical", []) or []
        languages = skills.get("languages", []) or []
        other = skills.get("other", []) or []
        if technical:
            story.append(Paragraph(escape("Technical Skills: " + ", ".join(technical)), body_style))
        if languages:
            story.append(Paragraph(escape("Languages: " + ", ".join(languages)), body_style))
        if other:
            story.append(Paragraph(escape(", ".join(other)), body_style))

    doc = SimpleDocTemplate(
        output_path,
        pagesize=LETTER,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    doc.build(story)


def main():
    parser = argparse.ArgumentParser(description="Generate resume DOCX and PDF from JSON.")
    parser.add_argument("--input", required=True, help="Path to resume JSON input.")
    parser.add_argument("--font-name", default="Calibri", help="Preferred font name.")
    parser.add_argument("--font-file", default="", help="Optional .ttf font file path.")
    parser.add_argument("--docx", default="", help="Output DOCX filename.")
    parser.add_argument("--pdf", default="", help="Output PDF filename.")
    args = parser.parse_args()

    data = load_json(args.input)
    base_name = os.path.splitext(os.path.basename(args.input))[0]
    docx_path = args.docx or (base_name + ".docx")
    pdf_path = args.pdf or (base_name + ".pdf")

    create_docx(data, docx_path, args.font_name)
    create_pdf(data, pdf_path, args.font_name, args.font_file)

    print("Wrote:", docx_path)
    print("Wrote:", pdf_path)


if __name__ == "__main__":
    main()
